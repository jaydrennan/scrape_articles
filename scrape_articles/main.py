from flask import Flask, request, render_template, send_file, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt
from pathlib import Path
from zipfile import ZipFile
import html
from weasyprint import HTML
from io import BytesIO
from typing import Dict, List
import newspaper
import logging

app = Flask(__name__)
app.template_folder = "templates"
app.static_folder = "static"

def is_unwanted_text(text: str) -> bool:
    """Check if text line should be filtered out."""
    # Skip empty lines
    if not text.strip():
        return True
    
    # Skip all-caps lines that are longer than 5 words (to preserve acronyms)
    if text.isupper() and len(text.split()) > 5:
        return True
    
    # Common phrases to filter out (case insensitive)
    unwanted_phrases = [
        "click here to sign up",
        "click here to get",
        "follow us on",
        "subscribe to",
        "sign up for our newsletter"
    ]
    
    return any(phrase in text.lower() for phrase in unwanted_phrases)

@app.route("/", methods=["GET"])
def home():
    return render_template("index.html")

@app.route("/process", methods=["POST"])
def process_urls():
    urls = request.form.get("urls", "")
    url_list = [url.strip() for url in urls.split('\n') if url.strip()]
    output_dir = Path("output")
    output_dir.mkdir(exist_ok=True)
    
    generated_files = []
    results = {
        "successful": [],
        "failed": [],
        "total_submitted": len(url_list),
        "total_completed": 0,
        "empty_content": []
    }
    
    for url in url_list:
        try:
            # Add timeout to prevent hanging
            article = newspaper.Article(url, timeout=10)
            article.download()
            article.parse()
            
            # Add logging for debugging
            if not article.text.strip():
                logging.warning(f"No content found for URL: {url}")
                results["empty_content"].append(url)
                results["failed"].append({"url": url, "error": "No content found"})
                continue

            # Create safe filename from title
            safe_title = "".join(c for c in article.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_title = safe_title[:50]  # Limit filename length
            
            # Create Word document for this article
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            
            # Add title
            title = doc.add_paragraph()
            title_run = title.add_run(article.title)
            title_run.bold = True
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(12)
            
            # Add paragraphs
            paragraphs = article.text.split('\n')
            for paragraph_text in paragraphs:
                if paragraph_text.strip() and not is_unwanted_text(paragraph_text):
                    text_paragraph = doc.add_paragraph()
                    text_run = text_paragraph.add_run(paragraph_text)
                    text_run.font.name = 'Times New Roman'
                    text_run.font.size = Pt(12)
            
            # Create HTML content for this article
            html_content = [
                '<html><body>',
                f'<h2>{html.escape(article.title)}</h2>'
            ]
            for paragraph_text in paragraphs:
                if paragraph_text.strip() and not is_unwanted_text(paragraph_text):
                    html_content.append(f'<p>{html.escape(paragraph_text)}</p>')
            html_content.append('</body></html>')
            
            # Save individual files
            docx_path = output_dir / f"{safe_title}.docx"
            pdf_path = output_dir / f"{safe_title}.pdf"
            
            # Save Word document
            doc.save(docx_path)
            
            # Generate PDF
            HTML(string=''.join(html_content)).write_pdf(pdf_path)
            
            results["successful"].append(url)
            results["total_completed"] += 1
            generated_files.extend([docx_path, pdf_path])
            
        except newspaper.article.ArticleException as e:
            logging.error(f"Article extraction error for {url}: {str(e)}")
            results["failed"].append({"url": url, "error": f"Article extraction error: {str(e)}"})
        except Exception as e:
            logging.error(f"Unexpected error processing {url}: {str(e)}")
            results["failed"].append({"url": url, "error": str(e)})
    
    # Create report file
    report_path = output_dir / "report.txt"
    with open(report_path, "w") as f:
        f.write(f"Processing Report\n{'='*50}\n\n")
        f.write(f"Total URLs submitted: {results['total_submitted']}\n")
        f.write(f"Successfully processed: {results['total_completed']}\n")
        f.write(f"Failed: {len(results['failed'])}\n")
        f.write(f"Empty content: {len(results['empty_content'])}\n\n")
        
        f.write("Successful URLs:\n")
        for url in results["successful"]:
            f.write(f"✓ {url}\n")
        
        f.write("\nFailed URLs:\n")
        for failure in results["failed"]:
            f.write(f"✗ {failure['url']}: {failure['error']}\n")
    
    generated_files.append(report_path)
    
    # Create ZIP with all files including report
    zip_path = output_dir / "articles.zip"
    with ZipFile(zip_path, 'w') as zip_file:
        for file_path in generated_files:
            zip_file.write(file_path, arcname=file_path.name)
    
    # Return both the file download and the results
    response = {
        "zip_url": "/download",  # We'll create this endpoint
        "results": {
            "total_submitted": results["total_submitted"],
            "total_completed": results["total_completed"],
            "successful": results["successful"],
            "failed": results["failed"],
            "empty_content": results["empty_content"]
        }
    }
    
    # Store the zip path in session or temporary storage
    app.config['CURRENT_ZIP'] = str(zip_path)
    
    return jsonify(response)

@app.route("/download", methods=["GET"])
def download_zip():
    zip_path = Path(app.config.get('CURRENT_ZIP'))
    if zip_path.exists():
        return send_file(
            zip_path,
            mimetype="application/zip",
            as_attachment=True,
            download_name="articles.zip"
        )
    return "No zip file available", 404

if __name__ == "__main__":
    app.run(debug=True)