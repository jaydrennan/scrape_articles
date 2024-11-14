from fastapi import FastAPI, Request, Form
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
import newspaper
from docx import Document
from docx.shared import Pt
from pathlib import Path
from zipfile import ZipFile
import html
from weasyprint import HTML
from io import BytesIO

app = FastAPI()

# Set up templates and static files
templates = Jinja2Templates(directory="templates")
app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/")
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/process")
async def process_urls(request: Request, urls: str = Form(...)):
    # Split URLs by newline
    url_list = [url.strip() for url in urls.split('\n') if url.strip()]
    
    # Create output directory if it doesn't exist
    output_dir = Path("output")
    output_dir.mkdir(exist_ok=True)
    
    # Initialize html_content list with header
    html_content = ['<html><body>']
    
    # Create Word document
    doc = Document()
    
    # Set default font for Word
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    for url in url_list:
        try:
            # Parse article
            article = newspaper.Article(url)
            article.download()
            article.parse()
            
            # Add title with same font size, just bold
            title = doc.add_paragraph()
            title_run = title.add_run(article.title)
            title_run.bold = True
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(12)
            
            # Split the text into paragraphs and add each one separately
            paragraphs = article.text.split('\n')
            for paragraph_text in paragraphs:
                if paragraph_text.strip():  # Only add non-empty paragraphs
                    text_paragraph = doc.add_paragraph()
                    text_run = text_paragraph.add_run(paragraph_text)
                    text_run.font.name = 'Times New Roman'
                    text_run.font.size = Pt(12)
            
            doc.add_paragraph()  # Add extra space between articles
            
            # Add to HTML content with preserved spacing
            html_content.append(f'<h2>{html.escape(article.title)}</h2>')
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    html_content.append(f'<p>{html.escape(paragraph_text)}</p>')
            html_content.append('<br>')
            
        except Exception as e:
            error_msg = f"Error processing {url}: {str(e)}"
            doc.add_paragraph(error_msg)
            html_content.append(f'<p>{html.escape(error_msg)}</p>')
    
    html_content.append('</body></html>')
    
    # Save files
    docx_path = output_dir / "articles.docx"
    pdf_path = output_dir / "articles.pdf"
    zip_path = output_dir / "articles.zip"
    
    # Save Word document
    doc.save(docx_path)
    
    # Generate PDF
    HTML(string=''.join(html_content)).write_pdf(pdf_path)
    
    # Create ZIP with both files
    with ZipFile(zip_path, 'w') as zip_file:
        zip_file.write(docx_path, arcname=docx_path.name)
        zip_file.write(pdf_path, arcname=pdf_path.name)
    
    return FileResponse(
        zip_path,
        filename="articles.zip",
        media_type="application/zip"
    )